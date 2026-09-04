#!/usr/bin/env python3
"""Synthetic regression tests for DOCX folder replacement."""

from __future__ import annotations

from contextlib import redirect_stderr, redirect_stdout
from io import BytesIO, StringIO
import json
from pathlib import Path
import re
from tempfile import TemporaryDirectory
from typing import cast
import unittest
from unittest.mock import patch
from uuid import UUID
import xml.etree.ElementTree as ElementTree
from zipfile import ZIP_DEFLATED, ZipFile

from PIL import Image
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, Protection
from openpyxl.worksheet.table import Table
from pptx import Presentation
from pptx.enum.text import MSO_ANCHOR, MSO_AUTO_SIZE
from pptx.oxml import parse_xml
from pptx.oxml.ns import qn
from pptx.util import Inches, Pt
# pypdf does not publish PEP 561 metadata for its generic object model.
from pypdf import PdfReader, PdfWriter
from pypdf.generic import ArrayObject, ByteStringObject, ContentStream, DecodedStreamObject, DictionaryObject, FloatObject, NameObject, NumberObject, TextStringObject
# skia-python does not publish PEP 561 stubs; this is the native rendering boundary.
import skia  # type: ignore[import-not-found]

from pipeline.folder_replacement import (
    FolderReplacementResult,
    parse_include_patterns,
    replace_input_folder,
)
from pipeline.folder_replacement.processor import ProgressFactory, ProgressReporter
from pipeline.folder_replacement.docx import _docx_ocr_backgrounds, replace_docx_file
from pipeline.folder_replacement.pptx import _pptx_ocr_backgrounds
from pipeline.folder_replacement.pdf import (
    _PdfPaintSpan,
    _PdfShownText,
    _PdfVisualRegion,
    _PdfReplacementSerializationError,
    _pdf_apply_paint_span_bullet_overrides,
    _pdf_apply_legacy_bullet_override,
    _pdf_content_has_annotations,
    _pdf_decode_composite_bytes,
    _pdf_expansion_geometry_is_known,
    _pdf_fitted_region_operations,
    _pdf_is_candidate_bullet_error,
    _pdf_text_advance,
)
from pipeline.folder_replacement.xlsx import _replace_drawing
from pipeline.folder_replacement.docx import _validate_docx_embedded_fonts
from pipeline.bounded_text_layout import (
    BoundedTextBox,
    BoundedTextParagraph,
    BoundedTextRun,
    PortableTextUnsupportedError,
    noto_typefaces,
)
from pipeline.portable_bullet_overrides import LegacyBulletOverride
from pipeline.portable_fonts import static_noto_bytes
from pipeline.ocr import BoundingPolygon, OcrRequest, OcrResult, OcrText, PixelPoint
from pipeline.ocr.provider import LocalContractTestSkip
from pipeline.text_replacement import (
    TextReplacementProvider,
    TextReplacementRequest,
    TextReplacementResult,
)
from pipeline.text_replacement_plugins.character_mask import CharacterMaskProvider



from folder_replacement_test_support import (
    FONT_PATH,
    FolderReplacementTestCase,
    _CountingOcrProvider,
    _EmptyOcrProvider,
    _FailingOcrProvider,
    _FailingReplacementProvider,
    _LowConfidenceOcrProvider,
    _RecordedProgress,
    _RecordingReplacementProvider,
    _VectorOutlineOcrProvider,
    _synthetic_pdf_visual_region,
)

class FolderReplacementDocxTests(FolderReplacementTestCase):
    # Verifies FR-2026-09-03-03 and FR-2026-09-04-01.
    def test_replaces_reachable_chart_rich_text_and_embedded_workbook_caches(self) -> None:
        class _PrefixReplacementProvider:
            def replace(self, request: TextReplacementRequest) -> TextReplacementResult:
                return TextReplacementResult(
                    request.text if request.is_filename else f"translated-{request.text}", 1.0
                )

        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            workbook = Workbook()
            worksheet = workbook.active
            assert worksheet is not None
            worksheet.title = "Data"
            worksheet["A1"] = "One"
            worksheet["A2"] = "Two"
            worksheet["B1"] = "Unrelated supporting text"
            skipped_sheet = workbook.create_sheet("Skipped")
            skipped_sheet["A1001"] = "Untranslated supporting text"
            workbook_bytes = BytesIO()
            workbook.save(workbook_bytes)
            with ZipFile(BytesIO(workbook_bytes.getvalue())) as source_archive:
                embedded_parts = {
                    entry.filename: source_archive.read(entry.filename)
                    for entry in source_archive.infolist()
                }
            embedded_output = BytesIO()
            with ZipFile(embedded_output, "w", ZIP_DEFLATED) as embedded_archive:
                for name, data in embedded_parts.items():
                    embedded_archive.writestr(name, data)
                embedded_archive.writestr("custom/unrelated.bin", b"retained")
            source = input_root / "chart.docx"
            with ZipFile(source, "w", ZIP_DEFLATED) as archive:
                archive.writestr(
                    "[Content_Types].xml",
                    b'''<?xml version="1.0"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/charts/chart1.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/chart2.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chart+xml"/>
  <Override PartName="/word/charts/userShapes/drawing1.xml" ContentType="application/vnd.openxmlformats-officedocument.drawingml.chartshapes+xml"/>
  <Override PartName="/word/embeddings/chart.xlsx" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"/>
</Types>''',
                )
                archive.writestr(
                    "_rels/.rels",
                    b'''<?xml version="1.0"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>''',
                )
                archive.writestr(
                    "word/document.xml",
                    b'''<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p/></w:body></w:document>''',
                )
                archive.writestr(
                    "word/_rels/document.xml.rels",
                    b'''<?xml version="1.0"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rIdChart" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" Target="charts/chart1.xml"/>
</Relationships>''',
                )
                archive.writestr(
                    "word/charts/chart1.xml",
                    b'''<?xml version="1.0"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <c:chart><c:title><c:tx><c:rich><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr b="1"/><a:t>Title</a:t></a:r></a:p></c:rich></c:tx></c:title>
  <c:plotArea><c:barChart><c:ser><c:cat><c:strRef><c:f>Data!$A$1:$A$2</c:f><c:strCache><c:ptCount val="2"/><c:pt idx="0"><c:v>One</c:v></c:pt><c:pt idx="1"><c:v>Two</c:v></c:pt></c:strCache></c:strRef></c:cat><c:val><c:numRef><c:f>Data!$B$1:$B$2</c:f><c:numCache><c:pt idx="0"><c:v>10</c:v></c:pt></c:numCache></c:numRef></c:val></c:ser><c:ser><c:cat><c:strRef><c:f>Data!$A$1+1</c:f><c:strCache><c:pt idx="0"><c:v>Unsupported</c:v></c:pt></c:strCache></c:strRef></c:cat></c:ser></c:barChart></c:plotArea></c:chart>
  <c:externalData r:id="rIdWorkbook"/>
</c:chartSpace>''',
                )
                archive.writestr(
                    "word/charts/chart2.xml",
                    b'''<?xml version="1.0"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><c:chart><c:title><c:tx><c:rich><a:p><a:r><a:t>Unreachable</a:t></a:r></a:p></c:rich></c:tx></c:title></c:chart></c:chartSpace>''',
                )
                archive.writestr(
                    "word/charts/_rels/chart1.xml.rels",
                    b'''<?xml version="1.0"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rIdWorkbook" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/package" Target="../embeddings/chart.xlsx"/>
  <Relationship Id="rIdUserShape" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chartUserShapes" Target="userShapes/drawing1.xml"/>
</Relationships>''',
                )
                archive.writestr(
                    "word/charts/userShapes/drawing1.xml",
                    b'''<?xml version="1.0"?>
<cdr:userShapes xmlns:cdr="http://schemas.openxmlformats.org/drawingml/2006/chartDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><cdr:relSizeAnchor><cdr:sp><cdr:txBody><a:p><a:r><a:rPr i="1"/><a:t>Label</a:t></a:r></a:p></cdr:txBody></cdr:sp></cdr:relSizeAnchor></cdr:userShapes>''',
                )
                archive.writestr("word/embeddings/chart.xlsx", embedded_output.getvalue())

            result = self._run(
                input_root, output_root, _EmptyOcrProvider(), _PrefixReplacementProvider(),
                diagnostics_enabled=True,
            )

            self.assertEqual(1, result.processed_files)
            with ZipFile(source) as source_archive:
                source_relationships = source_archive.read("word/_rels/document.xml.rels")
            with ZipFile(output_root / "chart.docx") as archive:
                chart = ElementTree.fromstring(archive.read("word/charts/chart1.xml"))
                unreachable_chart = ElementTree.fromstring(archive.read("word/charts/chart2.xml"))
                user_shape = ElementTree.fromstring(archive.read("word/charts/userShapes/drawing1.xml"))
                translated_workbook = archive.read("word/embeddings/chart.xlsx")
                self.assertEqual(archive.read("word/_rels/document.xml.rels"), source_relationships)
            with ZipFile(BytesIO(translated_workbook)) as embedded_archive:
                self.assertEqual(b"retained", embedded_archive.read("custom/unrelated.bin"))
            chart_namespace = "{http://schemas.openxmlformats.org/drawingml/2006/chart}"
            drawing_namespace = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
            title = chart.find(f".//{drawing_namespace}t")
            unreachable_title = unreachable_chart.find(f".//{drawing_namespace}t")
            assert title is not None and unreachable_title is not None
            self.assertEqual("translated-Title", title.text)
            self.assertEqual("Unreachable", unreachable_title.text)
            self.assertEqual(
                ["translated-One", "translated-Two", "Unsupported"],
                [value.text for value in chart.findall(f".//{chart_namespace}strCache/{chart_namespace}pt/{chart_namespace}v")],
            )
            self.assertEqual("Data!$A$1:$A$2", chart.findtext(f".//{chart_namespace}strRef/{chart_namespace}f"))
            self.assertEqual("10", chart.findtext(f".//{chart_namespace}numCache/{chart_namespace}pt/{chart_namespace}v"))
            user_shape_label = user_shape.find(f".//{drawing_namespace}t")
            assert user_shape_label is not None
            self.assertEqual("translated-Label", user_shape_label.text)
            translated = load_workbook(BytesIO(translated_workbook))
            self.assertEqual("translated-One", translated["Data"]["A1"].value)
            self.assertEqual("translated-Two", translated["Data"]["A2"].value)
            self.assertEqual("translated-Unrelated supporting text", translated["Data"]["B1"].value)
            self.assertEqual("Untranslated supporting text", translated["Skipped"]["A1001"].value)
            Document(str(output_root / "chart.docx"))
            diagnostic = json.loads(
                (output_root / "chart.docx.diagnostics.json").read_text(encoding="utf-8")
            )
            self.assertIn(
                "docx_chart_string_cache_unresolved",
                [entry["reason_code"] for entry in diagnostic["entries"]],
            )
            self.assertEqual(
                ["Skipped"],
                [
                    entry["worksheet_name"]
                    for entry in diagnostic["entries"]
                    if entry["reason_code"] == "xlsx_fast_mode_worksheet_skipped"
                ],
            )

            class _NestedProgress:
                def __init__(self) -> None:
                    self.started: list[tuple[str, int | None, str]] = []
                    self.advanced: list[str] = []
                    self.cleared = 0

                def start_nested(self, name: str, total: int | None, unit: str = "stage") -> None:
                    self.started.append((name, total, unit))

                def advance_nested(self, label: str) -> None:
                    self.advanced.append(label)

                def clear_nested(self) -> None:
                    self.cleared += 1

            typeface = skia.Typeface.MakeFromFile(str(FONT_PATH))
            assert typeface is not None
            nested_progress = _NestedProgress()
            replace_docx_file(
                source,
                root / "nested-progress.docx",
                _EmptyOcrProvider(),
                _PrefixReplacementProvider(),
                "en",
                "en",
                typeface,
                lambda _label: None,
                nested_progress=nested_progress,
            )
            self.assertEqual(1, len(nested_progress.started))
            self.assertTrue(nested_progress.started[0][0].startswith("chart1.xml: chart.xlsx"))
            self.assertEqual(3, nested_progress.started[0][1])
            self.assertEqual("replacement request", nested_progress.started[0][2])
            self.assertEqual(
                [
                    "xl/worksheets/sheet1.xml replacement 1",
                    "xl/worksheets/sheet1.xml replacement 2",
                    "xl/worksheets/sheet1.xml replacement 3",
                ],
                nested_progress.advanced,
            )
            self.assertEqual(1, nested_progress.cleared)

    # Verifies FR-2026-08-27-01.
    def test_uses_a_direct_document_background_for_embedded_word_image_ocr(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            document = Path(temporary_directory) / "input.docx"
            with ZipFile(document, "w", ZIP_DEFLATED) as archive:
                archive.writestr(
                    "word/document.xml",
                    """<?xml version=\"1.0\"?>
                    <w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">
                      <w:background w:color=\"102030\"/>
                    </w:document>""",
                )
                archive.writestr("word/media/image1.png", b"synthetic")

            self.assertEqual({"word/media/image1.png": (16, 32, 48)}, _docx_ocr_backgrounds(document))

    # Verifies FR-2026-09-02-05.
    def test_flowing_word_text_uses_maximal_emphasis_translation_runs(self) -> None:
        class _MappedReplacementProvider:
            def __init__(self) -> None:
                self.requests: list[TextReplacementRequest] = []
                self.replacements = {
                    "abb": "Alpha", "c": "Beta", "d": "γ", "e": "123", "f": "Delta",
                    "g": "Hello,", "h": "World", "i": "漢", "j": "字",
                    "k ": "Left", "l": "Right", "keep": "Keep",
                }

            def replace(self, request: TextReplacementRequest) -> TextReplacementResult:
                self.requests.append(request)
                if request.is_filename:
                    return TextReplacementResult(request.text, 1.0)
                return TextReplacementResult(self.replacements[request.text], 1.0)

        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            source = input_root / "document.docx"
            with ZipFile(source, "w", ZIP_DEFLATED) as archive:
                archive.writestr(
                    "word/document.xml",
                    b'''<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:pPr><w:jc w:val="center"/></w:pPr>
      <w:r><w:rPr><w:i/></w:rPr><w:t>a</w:t></w:r>
      <w:r><w:rPr><w:b/></w:rPr><w:t>bb</w:t></w:r>
      <w:r><w:rPr><w:color w:val="FF0000"/></w:rPr><w:t>c</w:t></w:r>
      <w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:t>d</w:t></w:r>
      <w:r><w:rPr><w:vertAlign w:val="subscript"/></w:rPr><w:t>e</w:t></w:r>
      <w:r><w:rPr><w:strike/></w:rPr><w:t>f</w:t></w:r>
    </w:p>
    <w:p><w:r><w:t>g</w:t></w:r><w:r><w:rPr><w:color w:val="00FF00"/></w:rPr><w:t>h</w:t></w:r></w:p>
    <w:p><w:r><w:t>i</w:t></w:r><w:r><w:rPr><w:color w:val="0000FF"/></w:rPr><w:t>j</w:t></w:r></w:p>
    <w:p><w:r><w:t xml:space="preserve">k </w:t></w:r><w:r><w:rPr><w:color w:val="808080"/></w:rPr><w:t>l</w:t></w:r></w:p>
    <w:p><w:hyperlink w:anchor="bookmark"><w:r><w:t>keep</w:t></w:r></w:hyperlink></w:p>
  </w:body>
</w:document>''',
                )
            provider = _MappedReplacementProvider()

            self._run(input_root, output_root, _EmptyOcrProvider(), provider)

            with ZipFile(output_root / "document.docx") as archive:
                document = ElementTree.fromstring(archive.read("word/document.xml"))
                data = archive.read("word/document.xml")
            namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            paragraphs = document.findall(f".//{namespace}body/{namespace}p")
            self.assertEqual(["Alpha ", "Beta", "γ", "123 ", "Delta"], [
                "".join(text.text or "" for text in run.findall(f"{namespace}t"))
                for run in paragraphs[0].findall(f"{namespace}r")
            ])
            self.assertIsNotNone(paragraphs[0].find(f"{namespace}pPr/{namespace}jc"))
            first_properties = paragraphs[0].find(f"{namespace}r/{namespace}rPr")
            self.assertIsNotNone(first_properties)
            assert first_properties is not None
            self.assertIsNotNone(first_properties.find(f"{namespace}b"))
            colour = paragraphs[0].findall(f"{namespace}r")[1].find(f"{namespace}rPr/{namespace}color")
            superscript = paragraphs[0].findall(f"{namespace}r")[2].find(f"{namespace}rPr/{namespace}vertAlign")
            subscript = paragraphs[0].findall(f"{namespace}r")[3].find(f"{namespace}rPr/{namespace}vertAlign")
            self.assertIsNotNone(colour)
            self.assertIsNotNone(superscript)
            self.assertIsNotNone(subscript)
            assert colour is not None and superscript is not None and subscript is not None
            self.assertEqual("FF0000", colour.get(f"{namespace}val"))
            self.assertEqual("superscript", superscript.get(f"{namespace}val"))
            self.assertEqual("subscript", subscript.get(f"{namespace}val"))
            self.assertIsNotNone(paragraphs[0].findall(f"{namespace}r")[4].find(f"{namespace}rPr/{namespace}strike"))
            self.assertEqual(["Hello,", "World"], [item.text for item in paragraphs[1].iter(f"{namespace}t")])
            self.assertEqual(["漢", "字"], [item.text for item in paragraphs[2].iter(f"{namespace}t")])
            self.assertEqual(["Left", "Right"], [item.text for item in paragraphs[3].iter(f"{namespace}t")])
            self.assertEqual(["Keep"], [item.text for item in paragraphs[4].iter(f"{namespace}t")])
            self.assertNotIn(b"Noto Sans", data)
            self.assertFalse((output_root / "document.docx.diagnostics.json").exists())
            self.assertEqual(
                ["abb", "c", "d", "e", "f", "g", "h", "i", "j", "k ", "l", "keep"],
                [request.text for request in provider.requests if not request.is_filename],
            )

    # Verifies FR-2026-09-02-05 and FR-2026-09-02-06.
    def test_flowing_word_structure_fallback_translates_and_reports_debug_details(self) -> None:
        class _PrefixReplacementProvider:
            def __init__(self) -> None:
                self.requests: list[TextReplacementRequest] = []

            def replace(self, request: TextReplacementRequest) -> TextReplacementResult:
                self.requests.append(request)
                return TextReplacementResult(request.text if request.is_filename else f"translated-{request.text}", 1.0)

        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            with ZipFile(input_root / "document.docx", "w", ZIP_DEFLATED) as archive:
                archive.writestr(
                    "word/document.xml",
                    b'''<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>plain</w:t></w:r></w:p>
    <w:p><w:bookmarkStart w:id="0" w:name="bookmark"/><w:r><w:t>bookmark</w:t></w:r></w:p>
    <w:p><w:hyperlink w:anchor="bookmark"><w:r><w:t>hyperlink</w:t></w:r></w:hyperlink></w:p>
    <w:p><w:r><w:t>marker</w:t><w:lastRenderedPageBreak/></w:r></w:p>
  </w:body>
</w:document>''',
                )
            provider = _PrefixReplacementProvider()

            result = self._run(
                input_root, output_root, _EmptyOcrProvider(), provider, diagnostics_enabled=True
            )

            with ZipFile(output_root / "document.docx") as archive:
                document = ElementTree.fromstring(archive.read("word/document.xml"))
            namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            self.assertEqual(
                ["translated-plain", "translated-bookmark", "translated-hyperlink", "translated-marker"],
                [item.text for item in document.iter(f"{namespace}t")],
            )
            self.assertEqual(1, len(result.diagnostic_sidecars))
            diagnostic = json.loads((output_root / "document.docx.diagnostics.json").read_text("utf-8"))
            entries = diagnostic["entries"]
            self.assertEqual(3, len(entries))
            self.assertEqual(
                ["docx_paragraph_child_bookmarkStart"], entries[0]["structure_reasons"]
            )
            self.assertEqual(
                ["docx_paragraph_child_hyperlink"], entries[1]["structure_reasons"]
            )
            self.assertEqual(
                ["docx_run_child_lastRenderedPageBreak"], entries[2]["structure_reasons"]
            )
            self.assertEqual([1, 2, 3], [entry["location"]["paragraph_index"] for entry in entries])
            self.assertTrue(all(entry["kind"] == "fallback" for entry in entries))
            self.assertTrue(all(
                entry["reason_code"] == "docx_flowing_paragraph_structure_fallback"
                for entry in entries
            ))
            self.assertEqual(
                ["plain", "bookmark", "hyperlink", "marker"],
                [request.text for request in provider.requests if not request.is_filename],
            )

    # Verifies FR-2026-09-02-07.
    def test_word_caption_cross_references_remain_live_and_share_caption_translation(self) -> None:
        class _MappedReplacementProvider:
            def __init__(self) -> None:
                self.requests: list[TextReplacementRequest] = []
                self.replacements = {
                    "label 1": "Figure 1",
                    "label ": "Figure ",
                    "1": "1",
                    "after reference": "This is a description.",
                    " after reference": "This is a description.",
                    "cached missing": "Missing result",
                    "cached malformed": "Malformed result",
                    "cached ambiguous": "Ambiguous result",
                    "unrelated": "Unused",
                }

            def replace(self, request: TextReplacementRequest) -> TextReplacementResult:
                self.requests.append(request)
                if request.is_filename:
                    return TextReplacementResult(request.text, 1.0)
                return TextReplacementResult(self.replacements[request.text], 1.0)

        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            source = input_root / "document.docx"
            self._write_complete_docx(source)
            with ZipFile(source) as archive:
                parts = {entry.filename: archive.read(entry.filename) for entry in archive.infolist()}
            parts["word/document.xml"] = b'''<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    <w:bookmarkStart w:id="1" w:name="caption_figure"/>
    <w:p>
      <w:r><w:t>label </w:t></w:r>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve"> SEQ Figure \\* ARABIC </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:r><w:t>1</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
      <w:bookmarkEnd w:id="1"/>
      <w:r><w:t xml:space="preserve"> after reference</w:t></w:r>
    </w:p>
    <w:p>
      <w:hyperlink r:id="rIdCrossReference">
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> REF caption_figure \\h </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:t>label 1</w:t></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
      </w:hyperlink>
      <w:r><w:t>after reference</w:t></w:r>
    </w:p>
    <w:p>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve"> REF caption_figure \\h </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:r><w:t>label 1</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
      <w:r><w:t>after reference</w:t></w:r>
    </w:p>
    <w:p>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve"> REF missing_caption \\h </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>cached missing</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    </w:p>
    <w:p>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve"> REF malformed_caption \\h </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>cached malformed</w:t></w:r>
    </w:p>
    <w:p><w:bookmarkStart w:id="2" w:name="ambiguous_caption"/><w:r><w:t>unrelated</w:t></w:r><w:bookmarkEnd w:id="2"/></w:p>
    <w:p><w:bookmarkStart w:id="3" w:name="ambiguous_caption"/><w:r><w:t>unrelated</w:t></w:r><w:bookmarkEnd w:id="3"/></w:p>
    <w:p>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve"> REF ambiguous_caption \\h </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>cached ambiguous</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    </w:p>
    <w:p>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve"> PAGEREF caption_figure \\h </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>2</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    </w:p>
    <w:sectPr/>
  </w:body>
</w:document>'''
            parts["word/_rels/document.xml.rels"] = b'''<?xml version="1.0"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rIdCrossReference" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="https://example.invalid/caption" TargetMode="External"/>
</Relationships>'''
            with ZipFile(source, "w", ZIP_DEFLATED) as archive:
                for name, contents in parts.items():
                    archive.writestr(name, contents)
            provider = _MappedReplacementProvider()

            self._run(input_root, output_root, _EmptyOcrProvider(), provider, diagnostics_enabled=True)

            output = output_root / "document.docx"
            Document(str(output))
            with ZipFile(output) as archive:
                document = ElementTree.fromstring(archive.read("word/document.xml"))
                relationships = archive.read("word/_rels/document.xml.rels")
            namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            paragraphs = document.findall(f".//{namespace}body/{namespace}p")
            self.assertEqual(
                "Figure 1 This is a description.",
                "".join(item.text or "" for item in paragraphs[0].iter(f"{namespace}t")),
            )
            self.assertEqual(["Figure ", "1", " ", "This is a description."], [
                item.text for item in paragraphs[0].iter(f"{namespace}t")
            ])
            self.assertEqual(
                "Figure 1 This is a description.",
                "".join(item.text or "" for item in paragraphs[1].iter(f"{namespace}t")),
            )
            self.assertEqual(
                "Figure 1 This is a description.",
                "".join(item.text or "" for item in paragraphs[2].iter(f"{namespace}t")),
            )
            self.assertEqual("Missing result", next(paragraphs[3].iter(f"{namespace}t")).text)
            self.assertEqual("Malformed result", next(paragraphs[4].iter(f"{namespace}t")).text)
            self.assertEqual("Ambiguous result", next(paragraphs[7].iter(f"{namespace}t")).text)
            self.assertEqual("2", next(paragraphs[8].iter(f"{namespace}t")).text)
            self.assertEqual(3, len(document.findall(f".//{namespace}bookmarkStart")))
            self.assertEqual(3, len(document.findall(f".//{namespace}bookmarkEnd")))
            self.assertEqual(20, len(document.findall(f".//{namespace}fldChar")))
            self.assertEqual(7, len(document.findall(f".//{namespace}instrText")))
            self.assertEqual(1, len(document.findall(f".//{namespace}hyperlink")))
            self.assertIn(b'rIdCrossReference', relationships)
            self.assertEqual(
                1,
                [request.text for request in provider.requests if not request.is_filename].count("label 1"),
            )
            diagnostic = json.loads((output_root / "document.docx.diagnostics.json").read_text("utf-8"))
            cross_reference_entries = [
                entry
                for entry in diagnostic["entries"]
                if entry["container_kind"] == "docx_cross_reference"
            ]
            self.assertEqual(
                [
                    "docx_cross_reference_bookmark_missing",
                    "docx_cross_reference_unbalanced_field",
                    "docx_cross_reference_bookmark_ambiguous",
                ],
                [entry["reason_code"] for entry in cross_reference_entries],
            )
            self.assertEqual(["REF", "REF", "REF"], [
                entry["field_kind"] for entry in cross_reference_entries
            ])
            self.assertEqual(
                ["missing_caption", "malformed_caption", "ambiguous_caption"],
                [entry["bookmark_identity"] for entry in cross_reference_entries],
            )
            self.assertEqual(
                ["cached missing", "cached malformed", "cached ambiguous"],
                [entry["field_result_text"] for entry in cross_reference_entries],
            )

    # Verifies FR-2026-09-02-08.
    def test_word_field_boundaries_preserve_system_fields_and_safe_joiners(self) -> None:
        class _RecordingIdentityProvider:
            def __init__(self) -> None:
                self.requests: list[TextReplacementRequest] = []

            def replace(self, request: TextReplacementRequest) -> TextReplacementResult:
                self.requests.append(request)
                return TextReplacementResult(request.text, 1.0)

        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            source = input_root / "document.docx"
            self._write_complete_docx(source)
            with ZipFile(source) as archive:
                parts = {entry.filename: archive.read(entry.filename) for entry in archive.infolist()}
            page_fields = b'''<w:r><w:fldChar w:fldCharType="begin"/></w:r>
<w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>7</w:t></w:r>
<w:r><w:fldChar w:fldCharType="end"/></w:r>'''
            page_count_field = b'''<w:r><w:fldChar w:fldCharType="begin"/></w:r>
<w:r><w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>
<w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>22</w:t></w:r>
<w:r><w:fldChar w:fldCharType="end"/></w:r>'''
            parts["word/document.xml"] = b'''<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t xml:space="preserve">P. </w:t></w:r>''' + page_fields + b'''<w:r><w:t>/</w:t></w:r>''' + page_count_field + b'''</w:p>
    <w:p>''' + page_fields + b'''<w:r><w:t>after</w:t></w:r></w:p>
    <w:p><w:r><w:t>left</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> DOCVARIABLE sample </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>fallback</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r><w:r><w:t>right</w:t></w:r></w:p>
    <w:sectPr/>
  </w:body>
</w:document>'''
            parts["word/footer1.xml"] = b'''<?xml version="1.0"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p><w:r><w:t xml:space="preserve">P. </w:t></w:r>''' + page_fields + b'''<w:r><w:t>/</w:t></w:r>''' + page_count_field + b'''</w:p>
</w:ftr>'''
            with ZipFile(source, "w", ZIP_DEFLATED) as archive:
                for name, contents in parts.items():
                    archive.writestr(name, contents)
            provider = _RecordingIdentityProvider()

            self._run(input_root, output_root, _EmptyOcrProvider(), provider)

            output = output_root / "document.docx"
            Document(str(output))
            namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            with ZipFile(output) as archive:
                document = ElementTree.fromstring(archive.read("word/document.xml"))
                footer = ElementTree.fromstring(archive.read("word/footer1.xml"))
            paragraphs = document.findall(f".//{namespace}body/{namespace}p")
            self.assertEqual(["P. ", "7", "/", "22"], [
                item.text for item in paragraphs[0].iter(f"{namespace}t")
            ])
            self.assertEqual(["7", " ", "after"], [
                item.text for item in paragraphs[1].iter(f"{namespace}t")
            ])
            self.assertEqual(["P. ", "7", "/", "22"], [
                item.text for item in footer.iter(f"{namespace}t")
            ])
            self.assertEqual(["PAGE", "NUMPAGES"], [
                (item.text or "").strip() for item in footer.iter(f"{namespace}instrText")
            ])
            requested = [request.text for request in provider.requests if not request.is_filename]
            self.assertNotIn("7", requested)
            self.assertNotIn("22", requested)
            self.assertNotIn("leftright", requested)
            self.assertTrue({"left", "fallback", "right"}.issubset(requested))

    # Verifies FR-2026-09-02-09.
    def test_word_toc_result_retains_tabs_and_page_fields_without_generated_spaces(self) -> None:
        class _MappedReplacementProvider:
            def __init__(self) -> None:
                self.requests: list[TextReplacementRequest] = []

            def replace(self, request: TextReplacementRequest) -> TextReplacementResult:
                self.requests.append(request)
                replacements = {"entry one": "Entry One", "entry two": "Entry Two"}
                return TextReplacementResult(
                    request.text if request.is_filename else replacements[request.text], 1.0
                )

        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            source = input_root / "document.docx"
            self._write_complete_docx(source)
            with ZipFile(source) as archive:
                parts = {entry.filename: archive.read(entry.filename) for entry in archive.infolist()}
            parts["word/document.xml"] = b'''<?xml version="1.0"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:pPr><w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9000"/></w:tabs></w:pPr>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve"> TOC \\h </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:hyperlink w:anchor="toc_one"><w:r><w:t>entry one</w:t></w:r><w:r><w:tab/></w:r>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> PAGEREF toc_one \\h </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>5</w:t></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r></w:hyperlink>
    </w:p>
    <w:p><w:pPr><w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9000"/></w:tabs></w:pPr>
      <w:hyperlink w:anchor="toc_two"><w:r><w:t>entry two</w:t></w:r><w:r><w:tab/></w:r>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> PAGEREF toc_two \\h </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>8</w:t></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r></w:hyperlink>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    </w:p>
    <w:sectPr/>
  </w:body>
</w:document>'''
            with ZipFile(source, "w", ZIP_DEFLATED) as archive:
                for name, contents in parts.items():
                    archive.writestr(name, contents)
            provider = _MappedReplacementProvider()

            self._run(input_root, output_root, _EmptyOcrProvider(), provider)

            output = output_root / "document.docx"
            Document(str(output))
            namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            with ZipFile(output) as archive:
                document = ElementTree.fromstring(archive.read("word/document.xml"))
            paragraphs = document.findall(f".//{namespace}body/{namespace}p")
            self.assertEqual(["Entry One", "5"], [
                item.text for item in paragraphs[0].iter(f"{namespace}t")
            ])
            self.assertEqual(["Entry Two", "8"], [
                item.text for item in paragraphs[1].iter(f"{namespace}t")
            ])
            self.assertEqual(4, len(document.findall(f".//{namespace}tab")))
            self.assertEqual(["dot", "dot"], [
                item.get(f"{namespace}leader")
                for item in document.findall(f".//{namespace}pPr/{namespace}tabs/{namespace}tab")
            ])
            self.assertEqual(2, len(document.findall(f".//{namespace}hyperlink")))
            self.assertEqual(["TOC \\h", "PAGEREF toc_one \\h", "PAGEREF toc_two \\h"], [
                (item.text or "").strip() for item in document.iter(f"{namespace}instrText")
            ])
            requested = [request.text for request in provider.requests if not request.is_filename]
            self.assertEqual(["entry one", "entry two"], requested)

    # Verifies FR-2026-08-04-07 and FR-2026-08-22-03.
    def test_docx_basic_layout_fits_drawing_text_and_embeds_conformant_fonts(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            source = input_root / "document.docx"
            self._write_complete_docx(source)
            self._run(
                input_root, output_root, _EmptyOcrProvider(),
                _RecordingReplacementProvider(replacement_text="Long replacement text " * 20),
                document_text_layout="preserve-basic-layout",
            )
            output = output_root / "document.docx"
            Document(str(output))
            with ZipFile(output) as archive:
                data = archive.read("word/document.xml")
                font_table = ElementTree.fromstring(archive.read("word/fontTable.xml"))
                font_relationships = ElementTree.fromstring(
                    archive.read("word/_rels/fontTable.xml.rels")
                )
                content_types = ElementTree.fromstring(archive.read("[Content_Types].xml"))
                settings = ElementTree.fromstring(archive.read("word/settings.xml"))
                self._assert_word_run_property_order(ElementTree.fromstring(data))
                targets = {item.get("Id"): item.get("Target") for item in font_relationships}
                embedded_font_ids: set[str] = set()
                namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                relationships_namespace = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
                font_sources = {
                    "Noto Sans JP": "sans-serif",
                    "Noto Serif JP": "serif",
                    "Noto Sans Mono": "fixed-width",
                }
                for family, classification in font_sources.items():
                    entries = [
                        item for item in font_table.findall(f"{namespace}font")
                        if item.get(f"{namespace}name") == family
                    ]
                    self.assertEqual(1, len(entries))
                    if family != "Noto Sans JP":
                        expected_metadata = self._sfnt_font_metadata(
                            static_noto_bytes(classification, False)
                        )
                        metadata_names = [
                            child.tag.rsplit("}", 1)[-1]
                            for child in entries[0]
                        ]
                        self.assertEqual(
                            [
                                "panose1", "charset", "family", "pitch", "sig",
                                "embedRegular", "embedBold",
                            ],
                            metadata_names,
                        )
                        for child in entries[0][:5]:
                            self.assertEqual(
                                expected_metadata[child.tag.rsplit("}", 1)[-1]],
                                {
                                    key.rsplit("}", 1)[-1]: value
                                    for key, value in child.attrib.items()
                                },
                            )
                    for style in ("embedRegular", "embedBold"):
                        embedding = entries[0].find(f"{namespace}{style}")
                        self.assertIsNotNone(embedding)
                        assert embedding is not None
                        relationship_id = embedding.get(f"{relationships_namespace}id")
                        font_key = embedding.get(f"{namespace}fontKey")
                        self.assertIsNotNone(relationship_id)
                        self.assertIsNotNone(font_key)
                        assert relationship_id is not None and font_key is not None
                        self.assertNotIn(relationship_id, embedded_font_ids)
                        embedded_font_ids.add(relationship_id)
                        target = targets[relationship_id]
                        recovered = bytearray(archive.read(f"word/{target}"))
                        standard_key = UUID(font_key.strip("{}")).bytes[::-1]
                        for offset in range(min(32, len(recovered))):
                            recovered[offset] ^= standard_key[offset % 16]
                        self.assertEqual(
                            bytes(recovered),
                            static_noto_bytes(classification, style == "embedBold"),
                        )
                        self.assertIsNotNone(
                            skia.Typeface.MakeFromData(skia.Data.MakeWithCopy(bytes(recovered)))
                        )
                self.assertEqual(6, len(embedded_font_ids))
                self.assertEqual(
                    1,
                    sum(item.get("Extension") == "odttf" for item in content_types),
                )
                self.assertEqual(
                    1,
                    sum(item.get("PartName") == "/word/fontTable.xml" for item in content_types),
                )
                content_type_names = [item.tag.rsplit("}", 1)[-1] for item in content_types]
                first_override = content_type_names.index("Override")
                self.assertNotIn("Default", content_type_names[first_override:])
                setting_names = [item.tag.rsplit("}", 1)[-1] for item in settings]
                self.assertEqual(
                    [
                        "writeProtection", "zoom", "embedTrueTypeFonts",
                        "bordersDoNotSurroundHeader",
                    ],
                    setting_names,
                )
            self.assertIn(b"Noto Sans JP", data)
            self.assertIn(b"Long replacement text", data)
            generated_property_names = {
                child.tag.rsplit("}", 1)[-1]
                for properties in ElementTree.fromstring(data).findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent"
                    "//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
                )
                for child in properties
            }
            self.assertTrue({"b", "i", "u"}.issubset(generated_property_names))

            corrupted = root / "corrupted.docx"
            with ZipFile(output) as source_archive, ZipFile(corrupted, "w", ZIP_DEFLATED) as corrupted_archive:
                corrupted_part = "word/fonts/pipeline-sans-serif-regular.odttf"
                for entry in source_archive.infolist():
                    payload = source_archive.read(entry.filename)
                    if entry.filename == corrupted_part:
                        payload = b"bad font" + payload[8:]
                    corrupted_archive.writestr(entry, payload)
            with ZipFile(corrupted) as corrupted_archive:
                parts = {
                    entry.filename: corrupted_archive.read(entry.filename)
                    for entry in corrupted_archive.infolist()
                }
            with self.assertRaisesRegex(ValueError, "not a loadable OpenType font"):
                _validate_docx_embedded_fonts(parts)

            missing_setting_parts = parts.copy()
            settings = ElementTree.fromstring(missing_setting_parts["word/settings.xml"])
            setting = settings.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}embedTrueTypeFonts")
            self.assertIsNotNone(setting)
            assert setting is not None
            settings.remove(setting)
            missing_setting_parts["word/settings.xml"] = ElementTree.tostring(settings)
            with self.assertRaisesRegex(ValueError, "embedded-font setting is missing"):
                _validate_docx_embedded_fonts(missing_setting_parts)

            misplaced_setting_parts = parts.copy()
            settings = ElementTree.fromstring(misplaced_setting_parts["word/settings.xml"])
            setting = settings.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}embedTrueTypeFonts")
            self.assertIsNotNone(setting)
            assert setting is not None
            settings.remove(setting)
            settings.append(setting)
            misplaced_setting_parts["word/settings.xml"] = ElementTree.tostring(settings)
            with self.assertRaisesRegex(ValueError, "not in CT_Settings order"):
                _validate_docx_embedded_fonts(misplaced_setting_parts)

    # Verifies FR-2026-08-22-11.
    def test_docx_drawing_text_resolves_default_paragraph_style_font_size(self) -> None:
        def text_box_font_sizes(path: Path) -> list[int]:
            with ZipFile(path) as archive:
                document = ElementTree.fromstring(archive.read("word/document.xml"))
            namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            return [
                int(size.get(f"{namespace}val", "0"))
                for size in document.findall(
                    f".//{namespace}txbxContent//{namespace}rPr/{namespace}sz"
                )
            ]

        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            short_input = root / "short-input"
            short_output = root / "short-output"
            short_input.mkdir()
            short_source = short_input / "document.docx"
            self._write_complete_docx(short_source)
            with ZipFile(short_source) as archive:
                entries = [
                    (entry, archive.read(entry.filename)) for entry in archive.infolist()
                ]
            widened_entries = [
                (
                    entry,
                    data.replace(
                        b'cx="914400" cy="457200"', b'cx="5486400" cy="5486400"'
                    ).replace(b'<w:sz w:val="48"/>', b'')
                    if entry.filename == "word/document.xml" else data,
                )
                for entry, data in entries
            ]
            with ZipFile(short_source, "w", ZIP_DEFLATED) as archive:
                for entry, data in widened_entries:
                    archive.writestr(entry, data)
                archive.writestr(
                    "word/styles.xml",
                    b'''<?xml version="1.0" encoding="UTF-8"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:rPr><w:sz w:val="21"/></w:rPr>
  </w:style>
</w:styles>''',
                )
            self._run(
                short_input, short_output, _EmptyOcrProvider(),
                _RecordingReplacementProvider(replacement_text="Short"),
                document_text_layout="preserve-basic-layout-source-font",
            )
            self.assertEqual(
                [21, 21, 21], text_box_font_sizes(short_output / "document.docx")
            )

    # Verifies FR-2026-08-28-03.
    def test_word_failure_records_safe_native_text_context_and_continues(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            with ZipFile(input_root / "document.docx", "w", ZIP_DEFLATED) as archive:
                archive.writestr(
                    "word/document.xml",
                    """<?xml version="1.0"?>
                    <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                      <w:body><w:p><w:r><w:t>synthetic request text</w:t></w:r></w:p></w:body>
                    </w:document>""",
                )
            self._write_png(input_root / "good.png")

            result = self._run(
                input_root,
                output_root,
                _EmptyOcrProvider(),
                _FailingReplacementProvider(),
                diagnostics_enabled=True,
            )

            self.assertEqual(1, result.failed_files)
            self.assertEqual(1, result.processed_files)
            self.assertFalse((output_root / "document.docx").exists())
            self.assertTrue((output_root / "good.png").exists())
            diagnostic = json.loads(
                (output_root / "document.docx.diagnostics.json").read_text(encoding="utf-8")
            )
            entry = diagnostic["entries"][0]
            self.assertEqual("RuntimeError", entry["exception_type"])
            self.assertEqual(["ValueError"], entry["exception_cause_types"])
            self.assertEqual(
                {
                    "stage": "native_xml",
                    "container_kind": "office_xml",
                    "operation": "text_replacement",
                    "location": {"package_part": "word/document.xml"},
                    "request": {
                        "kind": "text_replacement",
                        "source_language": "en",
                        "target_language": "en",
                        "is_filename": False,
                        "input_character_count": 22,
                    },
                },
                entry["failure_context"],
            )
            serialized = json.dumps(diagnostic)
            self.assertNotIn("synthetic request text", serialized)
            self.assertNotIn("Synthetic chained detail", serialized)

    # Verifies FR-2026-08-28-03.
    def test_word_embedded_image_failure_records_safe_ocr_context(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            input_root = root / "input"
            output_root = root / "output"
            input_root.mkdir()
            image_bytes = BytesIO()
            Image.new("RGB", (7, 5), (16, 32, 48)).save(image_bytes, format="PNG")
            with ZipFile(input_root / "document.docx", "w", ZIP_DEFLATED) as archive:
                archive.writestr(
                    "word/document.xml",
                    """<?xml version="1.0"?>
                    <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>""",
                )
                archive.writestr("word/media/image1.png", image_bytes.getvalue())

            result = self._run(
                input_root,
                output_root,
                _FailingOcrProvider(),
                _RecordingReplacementProvider(),
                diagnostics_enabled=True,
            )

            self.assertEqual(1, result.failed_files)
            entry = json.loads(
                (output_root / "document.docx.diagnostics.json").read_text(encoding="utf-8")
            )["entries"][0]
            self.assertEqual("RuntimeError", entry["exception_type"])
            self.assertEqual(["ValueError"], entry["exception_cause_types"])
            self.assertEqual(
                {
                    "stage": "embedded_bitmap",
                    "container_kind": "embedded_bitmap",
                    "operation": "ocr",
                    "location": {
                        "package_part": "word/media/image1.png",
                        "item_index": 1,
                    },
                    "request": {
                        "kind": "ocr",
                        "language": "en",
                        "image_width": 7,
                        "image_height": 5,
                        "image_mode": "RGB",
                    },
                },
                entry["failure_context"],
            )
            self.assertNotIn("Synthetic OCR chained detail", json.dumps(entry))



if __name__ == "__main__":
    unittest.main()
